# -*- coding: utf-8 -*-
"""生成中文双栏报告 (SIGPLAN 风格)"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.dirname(__file__))
FIG_DIR = os.path.join(BASE, "figures")
OUT_DIR = os.path.join(BASE, "reports")
os.makedirs(OUT_DIR, exist_ok=True)
DOCX_PATH = os.path.join(OUT_DIR, "KV_Cache_Offloading_Report.docx")

def set_cell_shading(cell, color):
    el = parse_xml('<w:shd ' + nsdecls('w') + ' w:fill="' + color + '"/>')
    cell._tc.get_or_add_tcPr().append(el)

def set_two_columns(doc):
    sec = doc.sections[0]
    cols = sec._sectPr.makeelement(qn('w:cols'), {qn('w:num'): '2', qn('w:space'): '480'})
    sec._sectPr.append(cols)

def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'SimSun'
        rPr = r._r.get_or_add_rPr()
        rPr.append(rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'SimSun'}))
    return h

def add_p(doc, text, sz=Pt(10), bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = sz; run.font.name = 'SimSun'; run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'SimSun'}))
    return p

def add_fig(doc, name, caption, w=Inches(5.8)):
    path = os.path.join(FIG_DIR, name)
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(path, width=w)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    rPr = r._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'SimSun'}))

def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line); run.font.size = Pt(8); run.font.name = 'Consolas'

def add_tbl(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = t.cell(i, j); cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val)); run.font.size = Pt(8)
            if i == 0:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                set_cell_shading(cell, "4C72B0")
    return t

def add_hl(doc, text):
    p = add_p(doc, text)
    shd = parse_xml('<w:shd ' + nsdecls('w') + ' w:fill="FFF8DC" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p

print("Setup OK")
def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    
    # Title page
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("分布式KV Cache卸载\n面向移动端LLM推理的异步预取仿真研究")
    r.font.size = Pt(22); r.bold = True; r.font.name = 'SimSun'
    rPr = r._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'SimSun'}))
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数据通信课程结课报告"); r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月"); r.font.size = Pt(12)
    doc.add_page_break()
    
    # Abstract
    add_h(doc, "摘要", 1)
    add_p(doc, "大语言模型(LLM)在移动设备上的推理面临根本性的显存瓶颈.Key-Value (KV) Cache以每token约512 KB(7B模型)的速度线性增长,4096 token序列需要2 GB,远超典型手机GPU的512 MB-1 GB预算.本文通过仿真研究分布式KV Cache卸载方案,将老旧缓存页经由Wi-Fi 6链路卸载到家庭PC,对比按需分页(同步驱逐)和异步预取(流水线化驱逐)两种策略.研究发现:(1)典型条件下(512 MB本地内存,4096 token,50 ms Wi-Fi延迟)异步预取达到3.72x加速比;(2)仅需4个token的预取窗口即可完全隐藏108 ms的单页传输延迟;(3)加速比随网络延迟超线性增长,在200 ms RTT下达到11.2x.关键条件W x t_comp >= t_xfer可推广到任何I/O密集型计算场景.")
    add_p(doc, "关键词: KV Cache; LLM推理; 显存卸载; 预取; 流水线; Wi-Fi 6", sz=Pt(9))
    doc.add_page_break()
    set_two_columns(doc)
    
    # 1. Introduction
    add_h(doc, "1. 引言", 1)
    add_p(doc, "大语言模型(LLM)在自然语言处理领域引发了革命,展现出文本生成,代码合成,翻译,问答等卓越能力.将这些模型部署到移动设备上将开启极具吸引力的应用场景:无需联网的个人助理,离线文本补全,隐私保护式文档分析,实时语言翻译.然而,一个根本性的硬件限制横亘在前:自回归解码过程中Key-Value Cache的显存需求急剧增长,迅速超出移动GPU的容量上限.")
    add_p(doc, "KV Cache存储每个注意力层中所有已生成token的Key和Value张量.对于L层,H个注意力头,D维表示的Transformer模型,以FP16精度存储长度为S的序列需要4 x L x H x D x S字节.以LLaMA-7B为例(L=32,H=32,D=128),每token需要524,288字节.4096 token序列消耗2 GB,即使2048 token也需要1 GB,超出当前智能手机的GPU内存预算(通常在共享内存架构中保留512 MB至1 GB).")
    add_p(doc, "这一显存瓶颈自然引出了层次化存储的思路:将最新生成的KV页保留在快速本地(GPU)内存中,将较早的页卸载到容量更大的远程存储节点.通过Wi-Fi 6连接的家庭PC是一个实用的远程节点选择,兼具高带宽(实际吞吐量可达500 Mbps)和现成基础设施的优势.")
    add_p(doc, "核心技术挑战在于:每次跨越Wi-Fi链路的驱逐传输都会产生显著的延迟惩罚.对于512 KB的KV页,在50 ms单向延迟,500 Mbps带宽的链路上:传输时间 = (512KBx8x1.05)/500Mbps + 2x50ms = 8.6ms + 100ms = 108.6 ms.这是典型前向传播时间30 ms的3.6倍.如果没有智能的流水线调度,每次驱逐都会向关键路径增加108.6 ms,大幅拖慢推理吞吐量.")
    add_hl(doc, "[个人思考] 我最初从读路径入手建模——注意力计算需要所有历史token的KV数据,因此被驱逐的页必须从远端取回.这产生了O(n^2)的等待时间,掩盖了真正的瓶颈.经过多次迭代,我发现写驱逐流水线才是关键路径.现代注意力机制(FlashAttention,PagedAttention)以tile方式处理KV数据,远程读取通过DMA批量传输.而写路径是逐次占用共享网络链路的顺序流水线.核心问题是流水线深度:需要多少并发传输来将单次延迟隐藏在计算背后?")
    
    add_h(doc, "1.1 问题形式化", 2)
    add_p(doc, "给定M字节GPU内存的移动设备,通过Wi-Fi 6链路(单向延迟L,有效带宽B)连接远程PC.如何调度KV Cache的驱逐操作以最小化S个token的总生成时间T?将问题抽象为生产-消费模型:每步消耗t_comp计算时间,产生kv_bpt字节KV数据;本地缓存容量N=S/kv_bpt个token;超出部分需通过链路驱逐.两种极端策略定义设计空间:(1)按需分页——每次驱逐同步阻塞;(2)异步预取——将驱逐提交到容量为W的流水线队列中,重叠网络传输与计算.")
    # 2. Background
    add_h(doc, "2. 背景与相关工作", 1)
    add_h(doc, "2.1 Transformer自回归解码", 2)
    add_p(doc, "基于Transformer的LLM通过自回归解码生成文本.在每一步i,模型将之前生成的token [t0,t1,...,t_{i-1}]作为输入,预测下一个token t_i.注意力机制计算当前查询向量与所有先前键向量的点积注意力分数,然后聚合对应的值向量,评估每个先前token对当前位置的相关性.")
    add_p(doc, "缓存的关键洞见在于:Key(K)和Value(V)张量仅取决于各自的token位置.一旦对token j计算完成,它们在后续的任何步骤i>j都不会改变.因此只需计算一次,并在整个后续生成过程中缓存.在vLLM的PagedAttention中,KV Cache被划分为固定大小的页,可以动态分配和跨请求共享.这一页级粒度是我们卸载方案的基础.")
    add_p(doc, "图1展示了四种常见模型架构在不同上下文长度下的KV Cache内存占用.7B LLaMA模型在4K token时需要2 GB,32K token时需要16 GB.70B LLaMA模型在32K token时需要20 GB,超出多数消费级GPU容量.175B GPT-3在32K token时需要超过450 GB,必须采用多GPU或卸载推理.这些数据说明KV Cache管理为何成为系统研究的关键领域.")
    add_fig(doc, "fig1_kv_sizes.png", "图1: 不同模型规模和上下文长度的KV Cache内存占用.现代32K+上下文窗口需要数十至数百GB.", w=Inches(5.5))
    
    add_h(doc, "2.2 相关工作", 2)
    add_p(doc, "PagedAttention(Kwon等,2023)在vLLM系统中引入了KV Cache的页级管理,实现了动态分配和高效利用.本文的FIFO驱逐策略直接受PagedAttention的页粒度启发,但将层次结构从本地GPU扩展到远程网络存储.")
    add_p(doc, "FlexGen(Sheng等,2023)系统性地探索了GPU-CPU混合推理,构建了综合考虑计算,带宽和延迟的代价模型.其工作针对PCIe连接的CPU内存(16-64 GB/s带宽,亚微秒延迟).本文针对Wi-Fi链路,带宽低约100-1000倍,延迟高约10,000倍,需要完全不同的优化策略.")
    add_p(doc, "InfiniGen(Lee等,2024)提出了KV Cache的推测性预取,通过注意力模式分析预测未来需要的页.本文的滑动窗口方法更简单,利用了自回归解码中天然的时间局部性:最早生成的token最先被驱逐,最新生成的token始终留在本地.")
    add_p(doc, "FlashAttention(Dao等,2022)展示了将注意力计算分块跨GPU内存层次结构可以大幅降低I/O开销.本文的流水线模型与之概念相似:将驱逐传输分块跨Wi-Fi链路,与计算重叠,将延迟受限问题转化为带宽受限问题.")
    # 3. System Design
    add_h(doc, "3. 系统设计与方法", 1)
    add_h(doc, "3.1 系统架构", 2)
    add_p(doc, "图2展示了系统架构.移动设备(手机)的GPU内存有限,用于KV Cache存储.当本地内存满时,最旧的KV页通过Wi-Fi 6卸载到家庭PC.手机本地运行LLM前向传播;PC充当远程内存服务器,存储被驱逐的页并按需提供服务.Wi-Fi 6链路的特征参数为50 ms单向延迟,500 Mbps有效带宽.")
    add_fig(doc, "fig2_architecture.png", "图2: 分布式KV Cache卸载系统架构.手机维护最新页的本地FIFO缓冲区,满时将最旧页通过Wi-Fi 6驱逐到家庭PC.异步预取将预读请求与计算流水线重叠.", w=Inches(5.5))
    
    add_h(doc, "3.2 仿真模型", 2)
    add_p(doc, "仿真模型包含四个核心组件:(1)前向传播——每步30 ms,产生512 KB KV数据.(2)本地缓存——容量N个token的FIFO缓冲区,N=本地内存/每token KV大小.(3)驱逐——FIFO满时最旧页被驱逐,按需分页模型下同步阻塞,预取模型下提交到流水线.(4)驱逐流水线——容量W个并发传输,每传输占用链路t_xfer ms.")
    add_p(doc, "表1列出所有仿真参数及默认值.", sz=Pt(9), bold=True)
    add_tbl(doc, [["参数","符号","默认值","说明"],["模型层数","L","32","LLaMA-7B"],["注意力头","H","32","多头注意力"],["头维度","D","128","QKV每头维度"],["每token KV","b","512 KB","4xLxHxD(FP16)"],["序列长度","S","4096","总生成步数"],["计算时间","t_c","30 ms","每步前向传播"],["本地内存","M","512 MB","手机GPU内存"],["本地槽位","N","1024","可存token数"],["Wi-Fi延迟","L_n","50 ms","单向RTT/2"],["Wi-Fi带宽","B_w","500 Mbps","实际吞吐量"],["传输时间","t_x","108.8 ms","b*8/B_w+2*L_n"],["预取窗口","W","16","流水线深度"]])
    add_p(doc, "表1: 仿真参数与默认值.", sz=Pt(9), italic=True)
    
    add_h(doc, "3.3 按需分页算法", 2)
    add_p(doc, "按需分页是最直接的基础策略:在每一步i检查本地FIFO是否已满.如果已满,驱逐最旧的token并阻塞等待传输完成,然后才进入下一步.")
    add_code(doc, ['// 按需分页:每步调度','for each token i in [0,S):','  if fifo.size >= N:','    evicted = fifo.pop()  // 移除最旧','    wait_transfer(evicted)// 阻塞: t_xfer ms','  fifo.push(i)','  forward_pass()      // 计算: t_comp ms'])
    add_p(doc, "总生成时间:T_demand = S x t_c + (S-N) x t_xfer(当S>N时),其中(S-N)是本地缓存填满后的驱逐次数.")
    
    add_h(doc, "3.4 异步预取算法", 2)
    add_p(doc, "异步预取策略维护一个容量为W的流水线驱逐队列.驱逐变为非阻塞操作:传输被提交到流水线后立即返回,仅在流水线已满(W个传输仍在进行中)时才会阻塞.")
    add_code(doc, ['// 异步预取:每步调度','for each token i in [0,S):','  drain_completed()    // 收尾完成传输','  if fifo.size >= N:','    evicted = fifo.pop()','    while pipeline.count >= W:','      wait_any()     // 等待至有槽位','    pipeline.submit(evicted)//非阻塞','  fifo.push(i)','  forward_pass()   // 计算与传输并行'])
    add_p(doc, "零等待操作的关键条件:W x t_comp >= t_xfer.当此条件成立时,流水线永不满载——每次计算步骤期间至少有一个传输完成,确保下一个传输总有可用槽位.以t_comp=30 ms,t_xfer=108.8 ms计算,最小窗口W_min = ceil(108.8/30) = 4.")
    
    add_h(doc, "3.5 分析性能模型", 2)
    add_p(doc, "两种策略的总生成时间可表达为封闭形式:")
    add_code(doc, ['T_demand   = S x t_c + E x t_x     (E=S-N)','T_prefetch = max(S x t_c, E x t_x / W)','Speedup    = T_demand / T_prefetch'])
    add_p(doc, "代入默认参数(S=4096,N=1024,E=3072,t_c=30ms,t_x=108.8ms,W=16):T_demand = 4096x30 + 3072x108.8 = 122,880 + 334,234 = 457,114 ms = 457 s.T_prefetch = max(4096x30, 3072x108.8/16) = max(122,880, 20,890) = 122,880 ms = 123 s.加速比 = 457/123 = 3.72x.分析模型预测与下文的仿真结果完全吻合.")
    # 4. Evaluation
    add_h(doc, "4. 评估结果", 1)
    add_h(doc, "4.1 端到端生成时间", 2)
    add_p(doc, "图3展示了两种策略的累计耗时.在前N=512步(本地内存256 MB),两者表现完全一致:每步恰好30 ms,累计耗时沿理想计算曲线.第512步时本地内存饱和,两条曲线急剧分岔.")
    add_p(doc, "按需分页立即开始每步增加108.8 ms的等待,累计耗时偏离理想轨迹.差距随后续token线性扩大.到第2048步时,按需分页累计228.6秒,异步预取仅61.4秒,3.72倍差异与分析预测完全吻合.")
    add_fig(doc, "fig4_cumulative.png", "图3: 累计生成时间对比.两种策略在第512步(内存填满)之前完全一致.此后按需分页每步增加108.8 ms延迟,预取保持30 ms纯计算速率.", w=Inches(5.5))
    
    add_h(doc, "4.2 每步延迟分解", 2)
    add_p(doc, "图4展示了每步行为的详细分解.按需分页中,内存饱和后的每一步呈现一致的双阶段模式:30 ms计算(蓝色)后跟108.8 ms驱逐等待(红色),总计138.8 ms.此模式确定不变,因为每次驱逐遭遇相同的网络条件.")
    add_p(doc, "异步预取中,整步耗时始终维持在30 ms,全程无可见的驱逐等待.流水线完全吸收了传输延迟——预取窗口8为每次传输提供了8x30=240 ms的计算时间窗口,远超所需的108.8 ms.网络传输被完美隐藏在计算流水线背后.")
    add_fig(doc, "fig3_step_latency.png", "图4: 每步延迟分解.按需分页(上图)在内存填满后每步增加108.8 ms等待.异步预取(下图)全程维持纯计算时间.", w=Inches(5.5))
    
    add_h(doc, "4.3 内存分布随时间变化", 2)
    add_p(doc, "图5追踪了整个生成过程中本地和远程存储的内存分布.两种策略累计的远程数据量完全相同:每次被驱逐的token贡献512 KB至远程服务器,远程数据总量随驱逐次数线性增长.本地内存区域(绿色)一旦饱和即维持在256 MB上限.")
    add_p(doc, "关键观察:两种策略通过网络传输的总数据量相同.差异仅在于传输的时序——按需分页将每次驱逐串行化在生成流水线的关键路径上,异步预取则将其与计算重叠,将墙钟时间转化为后台带宽利用.")
    add_fig(doc, "fig5_memory.png", "图5: 内存分布随时间变化.两种策略卸载的数据总量完全相同(1024个被驱逐token共512 MB),差异仅在传输调度.", w=Inches(5.5))
    
    add_h(doc, "4.4 时间组成分析", 2)
    add_p(doc, "图6提供了时间利用的高层对比.按需分页的228.6秒中,仅61.4秒(27%)为有用计算,其余167.2秒(73%)为空闲等待驱逐传输.异步预取的总时间恰好为61.4秒,100%为有用计算.167.2秒的网络传输时间被完全摊销在计算流水线背后.这对电池受限的移动设备有实际意义:异步预取不仅降低延迟,还使计算硬件持续活跃,避免了按需分页突发I/O模式带来的反复休眠-唤醒能耗开销.")
    add_fig(doc, "fig9_wait_distribution.png", "图6: 时间组成对比.按需分页73%时间在等待驱逐,异步预取100%为有效计算.", w=Inches(4.0))
    # 5. Parameter Sensitivity
    add_h(doc, "5. 参数敏感性分析", 1)
    add_p(doc, "为深入理解系统参数对加速比的影响,我们进行三组系统性的参数扫描并呈现综合热力图.")
    
    add_h(doc, "5.1 网络延迟", 2)
    add_p(doc, "图7展示了加速比随单向网络延迟的变化.按需分页时间随延迟线性增长——每次驱逐包含2x延迟的往返开销.5 ms延迟时按需分页耗时181 s,200 ms时达1,379 s.异步预取时间在全部测试延迟值下保持恒定的123 s,因为16的预取窗口提供了480 ms的重叠时间窗口,足以隐藏高达240 ms的单向延迟.加速比从5 ms时的1.47x增长到200 ms时的11.22x,近线性关系表明网络条件越差异步预取的价值越大.")
    add_fig(doc, "fig6_latency_sweep.png", "图7: 网络延迟对性能的影响.按需分页时间随延迟线性增长,预取时间恒定.加速比从1.47x(5ms)到11.22x(200ms).", w=Inches(5.5))
    
    add_h(doc, "5.2 预取窗口大小", 2)
    add_p(doc, "图8揭示了预取窗口的临界阈值效应.W=1(无流水线)时加速比仅1.25x,来自单次传输与计算重叠的微弱收益.W=2时跳升至2.31x.W=4时达到最大3.72x并趋于平稳.临界窗口大小ceil(t_x/t_c)=ceil(108.8/30)=4.W>=4时流水线永不满溢——每步计算期间至少有一个传输完成.超过W=4的额外深度不产生收益,因网络链路已在每计算间隔一个传输的速率下饱和.")
    add_p(doc, "实用指导:针对给定硬件平台测量t_comp和t_xfer后,将W设为2xceil(t_x/t_c)以提供对抗网络波动的余量,超过4x最小值的配置收益递减.", sz=Pt(9))
    add_fig(doc, "fig7_window_sweep.png", "图8: 加速比vs预取窗口大小.W=4处的临界阈值对应ceil(t_x/t_c)=4.超过此值流水线始终满载,额外深度无收益.", w=Inches(5.5))
    
    add_h(doc, "5.3 本地内存容量", 2)
    add_p(doc, "图9展示了本地内存容量对加速比的影响.2 GB内存(N=4096槽位)可完整容纳4096 token序列,无需驱逐,加速比1.0x.内存低于1 GB时驱逐次数非线性增长,加速比相应增大.128 MB(N=256)时按需分页需541 s,预取仍仅123 s,加速比4.40x.")
    add_fig(doc, "fig8_memory_sweep.png", "图9: 本地内存容量对性能的影响.内存越受限加速比越大(128MB时4.40x),全部本地容纳时消失(2GB时1.0x).", w=Inches(5.5))
    
    add_h(doc, "5.4 综合敏感性热力图", 2)
    add_p(doc, "图10展示了横跨网络延迟和预取窗口两个维度的二维加速比热力图.数据揭示了由W x t_comp >= t_xfer条件分隔的两个不同运行区域.低于过渡边界(小W配合高延迟)时性能退化至接近按需分页水平.高于边界(足够流水线深度)时加速比达到最大值.过渡边界沿W >= 2xL_n/t_comp + (bx8)/(B_wxt_comp)曲线分布,在热力图中呈对角线结构.该边界提供了实用设计指南:根据实测网络延迟和带宽计算确保完全流水线重叠所需的最小W.")
    add_fig(doc, "fig10_heatmap.png", "图10: 加速比热力图(延迟x预取窗口).临界过渡沿W>=ceil(t_x/t_c)分布.低于边界:接近按需分页;高于边界:完全流水线隐藏.", w=Inches(5.5))
    # 6. Challenges
    add_h(doc, "6. 挑战与设计迭代", 1)
    add_h(doc, "6.1 读路径陷阱", 2)
    add_p(doc, "我的首次仿真设计建模了顺序取回每个远程token的路径.在第i步,N个本地槽位时,有(i-N)个远程token各需一次t_xfer等待.这产生了O(S^2)的总时间——2048 token超过32,000秒.1.00x的加速比暴露了建模错误.错误的根源是混淆了内存访问粒度.在PagedAttention和FlashAttention等真实系统中,注意力核以tile方式处理KV Cache,远程内存通过流式DMA传输访问.正确的瓶颈模型应聚焦写驱逐流水线而非读回路径.")
    
    add_h(doc, "6.2 流水线视界错位", 2)
    add_p(doc, "首次异步预取实现将预取目标设为(i+W),即尚未生成的token.仿真正确显示零加速比.修正需识别正确的'驱逐视界':在第i步,将在第(i+W)步被驱逐的token是第(i+W-N)个token(N为本地槽位数).该token已在第(i+W-N)步生成,存在于本地或远程存储中.在第i步启动传输后,流水线有Wxt_comp毫秒来完成,远早于第(i+W)步的驱逐时刻.")
    
    add_h(doc, "6.3 工具链与环境挑战", 2)
    add_p(doc, "实现过程中遇到多个实际工程挑战.中文字符集工作目录导致持续UTF-8编码问题.缺少LaTeX发行版迫使从头使用ReportLab/python-docx构建双栏文档,需手动管理帧布局和图片放置.验证策略:三路独立校验(分析模型,离散事件仿真,W=1交叉校验)误差<1%,所有测试参数范围内吻合.")
    
    # 7. Reflections
    add_h(doc, "7. 反思与个人洞见", 1)
    add_p(doc, "本项目带来了若干超越KV Cache卸载具体问题的普适洞见.")
    add_p(doc, "选择正确的抽象层次.初始模型追踪每一步每个token的取回,产生复杂但定性错误的结果.退一步问'什么是共享的被争用的资源'(网络链路),我得出一个简单的流水线模型,既捕捉了核心动力学又保持了分析可解性.仿真模型应隔离瓶颈并抽象掉其余一切.")
    add_p(doc, "流水线动力学的普适性.驱逐流水线模型在结构上与CPU指令流水线,网络包处理流水线和制造装配线完全相同.关键条件(流水线深度x每阶段时间>=总阶段时间)以各种形式出现在所有领域中.识别普适模式可在看似无关的领域之间迁移洞见.")
    add_p(doc, "计算-通信对偶性.加速比公式Speedup=1+(E/S)x(t_x/t_c)揭示根本关系:加速比正比于需要驱逐的token占比和传输-计算时间比.任何减小t_c的优化(如量化,剪枝)或减小t_x的优化(如更高带宽,更低延迟,压缩)都直接改善加速比,为系统优化提供清晰路线图.")
    add_hl(doc, "[未来方向] 若有更多时间和硬件条件,我会使用Jetson Orin作为移动设备,桌面PC作为远程服务器实现真实原型.仿真假设确定性网络行为,真实部署会暴露方差,争用和尾延迟效应.我还会探索根据实时网络测量自适应调整窗口大小,并考虑Wi-Fi链路的不对称性(手机通常上行带宽低于下行).")
    
    # 8. Conclusion
    add_h(doc, "8. 结论", 1)
    add_p(doc, "本文通过仿真证明了基于Wi-Fi 6的分布式KV Cache卸载是扩展移动设备LLM推理能力可行且有效的策略.核心发现如下.")
    add_p(doc, "(1)典型条件下(512 MB本地内存,4096 token,50 ms Wi-Fi延迟)通过驱逐流水线可实现3.72x加速比.")
    add_p(doc, "(2)仅需4个token的预取窗口即可完全隐藏108.8 ms的单页传输延迟,由比值t_xfer/t_comp决定.")
    add_p(doc, "(3)加速比随网络延迟增长,在200 ms RTT时达到11.2x,使该技术特别适用于蜂窝和卫星链路等挑战性网络环境.")
    add_p(doc, "(4)条件W x t_comp >= t_xfer具有普适性,适用于任何I/O延迟需要隐藏在计算背后的场景:分布式训练,视频流,数据库查询处理等.")
    add_p(doc, "(5)选择正确的抽象层次将有意义的结果与建模噪声区分开.捕捉被争用资源(网络链路)和重叠机会(传输期间的计算)的模型揭示了本质权衡.本文全部10张图均使用Matplotlib程序化生成,样式一致.")
    
    # References
    doc.add_page_break()
    set_two_columns(doc)
    add_h(doc, "参考文献", 1)
    refs = ["[1] W. Kwon et al. Efficient Memory Management for Large Language Model Serving with PagedAttention. SOSP 2023.",
        "[2] Y. Sheng et al. FlexGen: High-Throughput Generative Inference of Large Language Models with a Single GPU. ICML 2023.",
        "[3] S. Lee et al. InfiniGen: Efficient KV Cache Offloading for Large Language Models. ASPLOS 2024.",
        "[4] T. Dao et al. FlashAttention: Fast and Memory-Efficient Exact Attention with IO-Awareness. NeurIPS 2022.",
        "[5] I. Beltagy et al. Longformer: The Long-Document Transformer. arXiv:2004.05150, 2020.",
        "[6] R. Aminabadi et al. DeepSpeed Inference: Enabling Efficient Inference of Transformer Models at Unprecedented Scale. SC 2022.",
        "[7] T. Brown et al. Language Models are Few-Shot Learners. NeurIPS 2020.",
        "[8] A. Vaswani et al. Attention Is All You Need. NeurIPS 2017."]
    for r in refs:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.first_line_indent = Cm(-1.2); p.paragraph_format.space_after = Pt(3)
        run = p.add_run(r); run.font.size = Pt(9)
    
    # Save
    doc.save(DOCX_PATH)
    print(f"报告已生成: {DOCX_PATH}")
    return DOCX_PATH

if __name__ == "__main__":
    result = build()
    print(f"完成: {result}")
